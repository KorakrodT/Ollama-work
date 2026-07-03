#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
สร้างหนังสือราชการไทยเป็นไฟล์ .docx ตามระเบียบงานสารบรรณ

วิธีใช้:
    python3 gen_official_doc.py spec.json output.docx
    cat spec.json | python3 gen_official_doc.py - output.docx

spec.json คือ JSON object ที่มีฟิลด์ "doc_type" และฟิลด์อื่น ๆ ตาม references/formats.md
รองรับ doc_type: external, internal, seal, command, regulation, bylaw,
               announcement, statement, certificate, minutes, project, speech
"""
import sys
import json
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEFAULT_FONT = "TH SarabunIT9"   # มาตรฐานราชการ (TH SarabunIT๙)
BODY_SIZE = 16
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
KRUT_PATH = os.path.join(SCRIPT_DIR, "..", "assets", "krut.png")


# ---------- low-level helpers ----------

def set_run_font(run, font_name, size=None, bold=False):
    run.font.name = font_name
    run.font.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    # ตั้งทั้ง ascii/hAnsi/cs ให้ไทยและละตินใช้ฟอนต์เดียวกัน
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), font_name)


def add_par(doc, text="", font=DEFAULT_FONT, size=BODY_SIZE, bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT, indent_cm=None, space_after=0,
            space_before=0, line_spacing=1.0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if indent_cm is not None:
        pf.first_line_indent = Cm(indent_cm)
    if text != "" or True:
        run = p.add_run(text)
        set_run_font(run, font, size, bold)
    return p


def add_runs(doc, segments, align=WD_ALIGN_PARAGRAPH.LEFT, font=DEFAULT_FONT,
             size=BODY_SIZE, indent_cm=None, space_after=0):
    """segments: list of (text, bold)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    if indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    for text, bold in segments:
        r = p.add_run(text)
        set_run_font(r, font, size, bold)
    return p


def setup_page(doc):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    # base style font
    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(BODY_SIZE)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), DEFAULT_FONT)


def add_krut(doc, height_cm=3.0, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    if os.path.exists(KRUT_PATH):
        run = p.add_run()
        run.add_picture(KRUT_PATH, height=Cm(height_cm))
    else:
        run = p.add_run("(ตราครุฑ สูง %.1f ซม.)" % height_cm)
        set_run_font(run, DEFAULT_FONT, 14)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p


def blank(doc, size=6):
    add_par(doc, "", size=size)


L = WD_ALIGN_PARAGRAPH.LEFT
C = WD_ALIGN_PARAGRAPH.CENTER
R = WD_ALIGN_PARAGRAPH.RIGHT
J = WD_ALIGN_PARAGRAPH.JUSTIFY


# ---------- signature block ----------

def add_signature(doc, sig, extra_top=True):
    if not sig:
        return
    if extra_top:
        blank(doc)
        blank(doc)
    name = sig.get("name", "(ลงชื่อ)")
    pos = sig.get("position", "")
    # ลงชื่อ จัดให้อยู่ค่อนไปทางขวา
    add_par(doc, name, align=C, indent_cm=None, space_after=0).paragraph_format.left_indent = Cm(8)
    if pos:
        add_par(doc, pos, align=C).paragraph_format.left_indent = Cm(8)


def add_signature_right(doc, sig):
    """ลายเซ็นแบบหนังสือภายนอก: ขอแสดงความนับถือ + ชื่อ + ตำแหน่ง กึ่งกลางฝั่งขวา"""
    if not sig:
        return
    name = sig.get("name", "(ลงชื่อ)")
    pos = sig.get("position", "")
    for txt in [name, pos]:
        if txt:
            p = add_par(doc, txt, align=C)
            p.paragraph_format.left_indent = Cm(8)


# ---------- document builders ----------

def build_external(doc, s):
    font = s.get("font", DEFAULT_FONT)
    add_krut(doc, 3.0)
    blank(doc)
    header = s.get("agency_header", [])
    # บรรทัดแรก = ที่ (ชิดซ้าย), ที่เหลือ = ชื่อ/ที่อยู่ (กึ่งกลางค่อนขวา)
    if header:
        add_par(doc, header[0], font=font, align=L, space_after=0)
        for line in header[1:]:
            p = add_par(doc, line, font=font, align=C, space_after=0)
            p.paragraph_format.left_indent = Cm(7)
    if s.get("date"):
        p = add_par(doc, s["date"], font=font, align=C, space_before=4, space_after=4)
        p.paragraph_format.left_indent = Cm(7)
    if s.get("subject"):
        add_runs(doc, [("เรื่อง  ", True), (s["subject"], False)], font=font, space_after=2)
    if s.get("salutation"):
        add_par(doc, s["salutation"], font=font, space_after=2)
    for ref in s.get("references", []):
        add_runs(doc, [("อ้างถึง  ", True), (ref, False)], font=font, space_after=0)
    encl = s.get("enclosures", [])
    if encl:
        if len(encl) == 1:
            add_runs(doc, [("สิ่งที่ส่งมาด้วย  ", True), (encl[0], False)], font=font, space_after=2)
        else:
            add_par(doc, "สิ่งที่ส่งมาด้วย", font=font, bold=True, space_after=0)
            for i, e in enumerate(encl, 1):
                add_par(doc, "%d. %s" % (i, e), font=font, indent_cm=1.5, space_after=0)
    blank(doc)
    for para in s.get("body", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4, line_spacing=1.0)
    blank(doc)
    closing = s.get("closing", "ขอแสดงความนับถือ")
    p = add_par(doc, closing, font=font, align=C)
    p.paragraph_format.left_indent = Cm(8)
    blank(doc)
    blank(doc)
    add_signature_right(doc, s.get("signature"))
    # contact
    c = s.get("contact")
    if c:
        blank(doc)
        if c.get("division"):
            add_par(doc, c["division"], font=font, size=14, space_after=0)
        line = []
        if c.get("phone"):
            line.append("โทร. " + c["phone"])
        if c.get("fax"):
            line.append("โทรสาร " + c["fax"])
        if line:
            add_par(doc, "  ".join(line), font=font, size=14, space_after=0)
    for cc in s.get("cc", []):
        add_par(doc, "สำเนาส่ง  " + cc, font=font, size=14, space_after=0)


def build_internal(doc, s):
    font = s.get("font", DEFAULT_FONT)
    # ครุฑเล็กมุมซ้าย + บันทึกข้อความ
    add_krut(doc, 1.5, align=L)
    add_par(doc, "บันทึกข้อความ", font=font, size=29, bold=True, align=C, space_after=6)
    seg = [("ส่วนราชการ  ", True), (s.get("agency", ""), False)]
    if s.get("phone"):
        seg.append(("  โทร. " + s["phone"], False))
    add_runs(doc, seg, font=font, space_after=0)
    add_runs(doc, [("ที่  ", True), (s.get("doc_no", ""), False),
                   ("          วันที่  ", True), (s.get("date", ""), False)],
             font=font, space_after=0)
    add_runs(doc, [("เรื่อง  ", True), (s.get("subject", ""), False)], font=font, space_after=4)
    add_par(doc, s.get("salutation", ""), font=font, space_after=4)
    for para in s.get("body", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4)
    add_signature(doc, s.get("signature"))


def build_seal(doc, s):
    font = s.get("font", DEFAULT_FONT)
    add_krut(doc, 3.0)
    blank(doc)
    for line in s.get("agency_header", []):
        add_par(doc, line, font=font, align=L, space_after=0)
    blank(doc)
    add_runs(doc, [("ถึง  ", True), (s.get("to", ""), False)], font=font, space_after=4)
    for para in s.get("body", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4)
    blank(doc)
    add_par(doc, s.get("agency_name", ""), font=font, align=C, space_after=0)
    add_par(doc, "(ประทับตราส่วนราชการ)", font=font, size=14, align=C, space_after=0)
    if s.get("date"):
        add_par(doc, s["date"], font=font, align=C, space_after=0)
    c = s.get("contact")
    if c:
        blank(doc)
        if c.get("division"):
            add_par(doc, c["division"], font=font, size=14, space_after=0)
        if c.get("phone"):
            add_par(doc, "โทร. " + c["phone"], font=font, size=14, space_after=0)


def _titled_directive(doc, s, head_word):
    """ใช้ร่วมกัน: คำสั่ง / ประกาศ / แถลงการณ์ (มีครุฑ + หัวกลาง)"""
    font = s.get("font", DEFAULT_FONT)
    add_krut(doc, 3.0)
    title = head_word + s.get("agency_name", "")
    add_par(doc, title, font=font, size=22, bold=True, align=C, space_after=0)
    if s.get("doc_no"):
        add_par(doc, s["doc_no"], font=font, align=C, space_after=0)
    if s.get("subject"):
        add_runs(doc, [("เรื่อง  ", True), (s["subject"], False)], font=font, align=C, space_after=4)
    add_par(doc, "_______________", font=font, align=C, space_after=4)
    return font


def build_command(doc, s):
    font = _titled_directive(doc, s, "คำสั่ง")
    for para in s.get("preamble", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4)
    for i, cl in enumerate(s.get("clauses", []), 1):
        if isinstance(cl, dict):
            no = cl.get("no", str(i))
            txt = cl.get("text", "")
        else:
            no, txt = str(i), cl
        add_par(doc, "ข้อ %s  %s" % (no, txt), font=font, align=J, indent_cm=2.5, space_after=4)
    if s.get("effective"):
        add_par(doc, s["effective"], font=font, align=J, indent_cm=2.5, space_after=4)
    blank(doc)
    if s.get("ordered_date"):
        p = add_par(doc, s["ordered_date"], font=font, align=C)
        p.paragraph_format.left_indent = Cm(6)
    blank(doc)
    add_signature_right(doc, s.get("signature"))


def build_regulation(doc, s, head="ระเบียบ"):
    font = s.get("font", DEFAULT_FONT)
    add_krut(doc, 3.0)
    add_par(doc, head + s.get("agency_name", ""), font=font, size=22, bold=True, align=C, space_after=0)
    title = "ว่าด้วย" + s.get("title", "")
    add_par(doc, title, font=font, size=18, bold=True, align=C, space_after=0)
    if s.get("year"):
        add_par(doc, "พ.ศ. " + s["year"], font=font, size=18, bold=True, align=C, space_after=4)
    add_par(doc, "_______________", font=font, align=C, space_after=4)
    for para in s.get("preamble", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4)
    for i, cl in enumerate(s.get("clauses", []), 1):
        if isinstance(cl, dict):
            no = cl.get("no", str(i))
            txt = cl.get("text", "")
        else:
            no, txt = str(i), cl
        add_par(doc, "ข้อ %s  %s" % (no, txt), font=font, align=J, indent_cm=2.5, space_after=4)
    blank(doc)
    if s.get("announced_date"):
        p = add_par(doc, s["announced_date"], font=font, align=C)
        p.paragraph_format.left_indent = Cm(6)
    blank(doc)
    add_signature_right(doc, s.get("signature"))


def build_bylaw(doc, s):
    build_regulation(doc, s, head="ข้อบังคับ")


def build_announcement(doc, s):
    font = _titled_directive(doc, s, "ประกาศ")
    for para in s.get("body", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4)
    blank(doc)
    if s.get("announced_date"):
        p = add_par(doc, s["announced_date"], font=font, align=C)
        p.paragraph_format.left_indent = Cm(6)
    blank(doc)
    add_signature_right(doc, s.get("signature"))


def build_statement(doc, s):
    font = _titled_directive(doc, s, "แถลงการณ์")
    for para in s.get("body", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4)
    blank(doc)
    add_par(doc, s.get("agency_name", ""), font=font, align=C, space_after=0)
    if s.get("announced_date"):
        add_par(doc, s["announced_date"], font=font, align=C, space_after=0)


def build_certificate(doc, s):
    font = s.get("font", DEFAULT_FONT)
    add_krut(doc, 3.0)
    blank(doc)
    if s.get("doc_no"):
        add_par(doc, s["doc_no"], font=font, align=C, space_after=4)
    add_par(doc, "หนังสือรับรอง", font=font, size=20, bold=True, align=C, space_after=4)
    for para in s.get("body", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4)
    blank(doc)
    if s.get("given_date"):
        p = add_par(doc, s["given_date"], font=font, align=C)
        p.paragraph_format.left_indent = Cm(6)
    blank(doc)
    add_signature_right(doc, s.get("signature"))


def build_minutes(doc, s):
    font = s.get("font", DEFAULT_FONT)
    add_par(doc, "รายงานการประชุม" + s.get("meeting_name", ""), font=font, size=18, bold=True, align=C, space_after=0)
    for t in s.get("time", []) if isinstance(s.get("time"), list) else ([s["time"]] if s.get("time") else []):
        add_par(doc, t, font=font, align=C, space_after=0)
    add_par(doc, "_______________", font=font, align=C, space_after=4)

    def name_list(label, items):
        if not items:
            return
        add_par(doc, label, font=font, bold=True, space_after=0)
        for i, it in enumerate(items, 1):
            add_par(doc, "%d. %s" % (i, it), font=font, indent_cm=1.0, space_after=0)

    name_list("ผู้มาประชุม", s.get("attendees", []))
    name_list("ผู้ไม่มาประชุม", s.get("absentees", []))
    name_list("ผู้เข้าร่วมประชุม", s.get("participants", []))
    if s.get("start_time"):
        add_par(doc, "เริ่มประชุมเวลา  " + s["start_time"], font=font, space_before=4, space_after=4)
    for ag in s.get("agenda", []):
        no = ag.get("no", "")
        title = ag.get("title", "")
        add_par(doc, "ระเบียบวาระที่ %s  %s" % (no, title), font=font, bold=True, space_after=2)
        for c in ag.get("content", []):
            add_par(doc, c, font=font, align=J, indent_cm=2.5, space_after=2)
    if s.get("end_time"):
        add_par(doc, "เลิกประชุมเวลา  " + s["end_time"], font=font, space_before=4, space_after=6)
    blank(doc)
    rec = s.get("recorder")
    chk = s.get("checker")
    if rec:
        add_par(doc, "(%s)" % rec.get("name", ""), font=font, align=C).paragraph_format.left_indent = Cm(8)
        add_par(doc, rec.get("position", "ผู้จดรายงานการประชุม"), font=font, align=C).paragraph_format.left_indent = Cm(8)
    if chk:
        blank(doc)
        add_par(doc, "(%s)" % chk.get("name", ""), font=font, align=C).paragraph_format.left_indent = Cm(8)
        add_par(doc, chk.get("position", "ผู้ตรวจรายงานการประชุม"), font=font, align=C).paragraph_format.left_indent = Cm(8)


def build_project(doc, s):
    font = s.get("font", DEFAULT_FONT)
    add_par(doc, "โครงการ" + s.get("project_name", ""), font=font, size=18, bold=True, align=C, space_after=6)

    def section(title, items):
        if not items:
            return
        add_par(doc, title, font=font, bold=True, space_after=0)
        if isinstance(items, str):
            add_par(doc, items, font=font, align=J, indent_cm=2.5, space_after=4)
        else:
            for i, it in enumerate(items, 1):
                add_par(doc, "%d. %s" % (i, it), font=font, align=J, indent_cm=1.0, space_after=0)
            blank(doc)

    section("๑. หลักการและเหตุผล", s.get("principle", []))
    section("๒. วัตถุประสงค์", s.get("objectives", []))
    section("๓. เป้าหมาย", s.get("targets", []))
    section("๔. วิธีดำเนินการ", s.get("methods", []))
    if s.get("duration"):
        add_par(doc, "๕. ระยะเวลาดำเนินการ  " + s["duration"], font=font, space_after=4)
    if s.get("place"):
        add_par(doc, "๖. สถานที่ดำเนินการ  " + s["place"], font=font, space_after=4)
    if s.get("budget"):
        add_par(doc, "๗. งบประมาณ  " + s["budget"], font=font, space_after=4)
    if s.get("responsible"):
        add_par(doc, "๘. ผู้รับผิดชอบโครงการ  " + s["responsible"], font=font, space_after=4)
    section("๙. การประเมินผล", s.get("evaluation", []))
    section("๑๐. ผลที่คาดว่าจะได้รับ", s.get("outcomes", []))


def build_speech(doc, s):
    font = s.get("font", DEFAULT_FONT)
    if s.get("title"):
        add_par(doc, s["title"], font=font, size=18, bold=True, align=C, space_after=6)
    for g in s.get("greeting", []):
        add_par(doc, g, font=font, align=C, space_after=0)
    blank(doc)
    for para in s.get("body", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4)
    for para in s.get("closing", []):
        add_par(doc, para, font=font, align=J, indent_cm=2.5, space_after=4)


BUILDERS = {
    "external": build_external,
    "internal": build_internal,
    "seal": build_seal,
    "command": build_command,
    "regulation": build_regulation,
    "bylaw": build_bylaw,
    "announcement": build_announcement,
    "statement": build_statement,
    "certificate": build_certificate,
    "minutes": build_minutes,
    "project": build_project,
    "speech": build_speech,
}


def main():
    if len(sys.argv) < 3:
        sys.exit("usage: gen_official_doc.py <spec.json|-> <output.docx>")
    spec_arg, out_path = sys.argv[1], sys.argv[2]
    raw = sys.stdin.read() if spec_arg == "-" else open(spec_arg, encoding="utf-8").read()
    spec = json.loads(raw)

    dtype = spec.get("doc_type")
    if dtype not in BUILDERS:
        sys.exit("unknown doc_type: %r (valid: %s)" % (dtype, ", ".join(BUILDERS)))

    # อนุญาตให้ override ฟอนต์มาตรฐานทั้งเอกสาร
    global DEFAULT_FONT
    if spec.get("font"):
        DEFAULT_FONT = spec["font"]

    doc = Document()
    setup_page(doc)
    BUILDERS[dtype](doc, spec)
    doc.save(out_path)
    print("saved:", out_path)


if __name__ == "__main__":
    main()
